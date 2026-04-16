"""
processors/text_processor.py
────────────────────────────────────────────────────────────────────
TXT / MD / DOCX / DOC → RAG chunk üretici.

Strateji:
  - TXT / MD → direkt oku, _chunk_text() ile böl
  - DOCX     → Hiyerarşik parçalama:
                  • Heading stillerini takip eder (H1, H2, H3)
                  • Her paragraf, ait olduğu başlık zincirini metadata
                    olarak ve chunk metninin başına "[Doküman > Bölüm > Alt Bölüm]"
                    ön eki olarak taşır.
                  • Tablolar "Data-to-Text" benzeri biçimde korunur.
  - DOC      → python-docx (eski format desteği kısıtlı)

İyileştirme (Faz 2 — Hiyerarşik Parçalama):
  Standart: "Bu oran geçen yıla göre %20 arttı."
  Yeni    : "[2023 İK Raporu > Maaş Zamları] Bu oran geçen yıla göre %20 arttı."
"""

import os
import uuid

_CHUNK_SIZE    = 1500
_CHUNK_OVERLAP = 200

# Word'deki başlık stil isimleri → hiyerarşi kademesi
_HEADING_STYLES = {
    "heading 1": 1, "başlık 1": 1, "title": 1,
    "heading 2": 2, "başlık 2": 2, "subtitle": 2,
    "heading 3": 3, "başlık 3": 3,
    "heading 4": 4, "başlık 4": 4,
}


def parse_text(
    file_path: str,
    original_name: str | None = None,
) -> list[dict]:
    file_basename = original_name or os.path.basename(file_path)
    ext = file_basename.rsplit(".", 1)[-1].lower() if "." in file_basename else ""

    try:
        if ext in ("docx", "doc"):
            return _parse_docx_hierarchical(file_path, file_basename, ext)
        elif ext == "pdf":
            return _parse_pdf(file_path, file_basename)
        else:
            raw_text = _read_plain(file_path)
            return _chunks_from_text(raw_text, file_basename, ext)
    except Exception as e:
        return [{
            "id":   f"error-{uuid.uuid4()}",
            "text": f"[{file_basename}] Dosya okunamadı: {e}",
            "metadata": {"source": file_basename, "error": str(e)}
        }]

# ── PDF ayrıştırıcı (Hiyerarşik PyMuPDF) ────────────────────────────────────────
def _parse_pdf(file_path: str, file_basename: str) -> list[dict]:
    """
    PDF'i basit dümdüz okumak yerine bloklara böler ve hiyerarşik yapı (başlık tahmini) yapar.
    Olası en iyi 'prev_id' ve 'next_id' akışını sağlar.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF kurulu değil. pip install PyMuPDF")

    doc = fitz.open(file_path)
    chunks = []
    
    current_block_lines = []
    current_prefix = f"[{file_basename}]"
    
    def _flush_pdf_block(page_num, total_pages):
        nonlocal current_block_lines, chunks
        if not current_block_lines:
            return
            
        block_text = "\n".join(current_block_lines)
        if not block_text.strip():
            return
            
        sub_chunks = _chunk_text(block_text)
        for idx, part in enumerate(sub_chunks):
            header = f"{current_prefix} [Sayfa {page_num}]\n"
            chunks.append({
                "id": str(uuid.uuid4()),
                "text": header + part,
                "metadata": {
                    "page": page_num,
                    "chunk_index": len(chunks) + 1,
                    "source": file_basename,
                    "type": "text_pdf",
                    "total_pages": total_pages,
                    "section": current_prefix
                }
            })
        current_block_lines = []

    total_pages = len(doc)
    
    for page_num, page in enumerate(doc, 1):
        # dict metoduyla metin stillerini yakalıyoruz
        blocks = page.get_text("dict").get("blocks", [])
        
        for b in blocks:
            if b.get("type") == 0:  # 0 text block
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                            
                        # Başlık karakter tespiti (Font boyutu > 13)
                        size = span.get("size", 10)
                        if size > 13 and len(text.split()) < 20: 
                            # Yeni bloğa geç (başlık değişti)
                            _flush_pdf_block(page_num, total_pages)
                            current_prefix = f"[{file_basename} > {text}]"
                        else:
                            current_block_lines.append(text)
                            
        # Sayfa geçişlerinde metni çok dağıtmamak adına ufak bir boşluk koyalım
        current_block_lines.append("\n")

    # Kalanı flush yap
    _flush_pdf_block(total_pages, total_pages)
    doc.close()
    
    # chunk'lara prev_id / next_id atama
    for i, chunk in enumerate(chunks):
        chunk["metadata"]["prev_id"] = chunks[i - 1]["id"] if i > 0 else ""
        chunk["metadata"]["next_id"] = chunks[i + 1]["id"] if i < len(chunks) - 1 else ""
        
    if not chunks:
        return [{
            "id": f"empty-{uuid.uuid4()}",
            "text": f"[{file_basename}] Dosya boş veya içerik bulunamadı.",
            "metadata": {"source": file_basename, "type": "text_empty"}
        }]
        
    return chunks


# ── Hiyerarşik DOCX ayrıştırıcı ──────────────────────────────────────

def _parse_docx_hierarchical(file_path: str, file_basename: str, ext: str) -> list[dict]:
    """
    DOCX/DOC'u başlık stillerini takip ederek hiyerarşik olarak parçalar.

    Algoritma:
      1. Belgedeki her paragrafı sırayla incele.
      2. Paragraf bir Heading ise → aktif başlık zincirini güncelle.
      3. Normal paragraf ise → aktif başlık zinciriyle birlikte bir
         "bağlamsal blok"a ekle.
      4. Blok _CHUNK_SIZE'ı aştığında yeni chunk oluştur (başlık her chunk'ta korunur).
    """
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx kurulu değil. pip install python-docx")

    doc    = docx.Document(file_path)
    chunks = []

    # Aktif başlık zinciri: {1: "Ana Başlık", 2: "Alt Başlık", 3: "..."}
    headings: dict[int, str] = {}

    # Mevcut blok — aynı hiyerarşi altındaki paragraflar birikir
    current_block_lines: list[str] = []
    current_prefix      = ""

    def _flush_block():
        """Birikmiş bloğu chunk'lara dönüştür."""
        nonlocal current_block_lines, current_prefix
        if not current_block_lines:
            return

        block_text = "\n\n".join(current_block_lines)
        sub_chunks = _chunk_text(block_text)
        total      = len(sub_chunks)

        for idx, part in enumerate(sub_chunks):
            # Metadata + metin ön eki
            header_comment = (
                f"{current_prefix}\n"
                if current_prefix else
                f"[{file_basename}]\n"
            )
            chunk_text_out = header_comment + part

            chunks.append({
                "id":   str(uuid.uuid4()),
                "text": chunk_text_out,
                "metadata": {
                    "page":          len(chunks) + 1,
                    "chunk_index":   idx + 1,
                    "total_in_block": total,
                    "source":        file_basename,
                    "type":          f"text_{ext}",
                    "section":       current_prefix,
                }
            })

        current_block_lines = []

    # ── Belge elemanlarını sırayla işle (paragraf + tablo birlikte) ──
    from docx.oxml.ns import qn

    doc_element = doc.element.body
    para_map    = {p._element: p for p in doc.paragraphs}
    table_map   = {t._element: t for t in doc.tables}

    for child in doc_element.iterchildren():
        # Paragraf
        if child.tag == qn("w:p"):
            para = para_map.get(child)
            if para is None:
                continue

            text  = para.text.strip()
            style = para.style.name.lower() if para.style else ""
            level = _HEADING_STYLES.get(style)

            if level is not None:
                # ── Başlık bulundu: bloğu kapat, zinciri güncelle ──
                _flush_block()

                # Daha derin seviyeleri sıfırla
                for k in list(headings.keys()):
                    if k >= level:
                        del headings[k]
                headings[level] = text

                # Yeni ön ek oluştur: [Belge > Bölüm 1 > Alt Bölüm A]
                chain_parts = [file_basename] + [
                    headings[k] for k in sorted(headings.keys()) if headings.get(k)
                ]
                current_prefix = "[" + " > ".join(chain_parts) + "]"

            else:
                # Düz paragraf → bloğa ekle (markdown formatında)
                if text:
                    current_block_lines.append(text)

        # Tablo
        elif child.tag == qn("w:tbl"):
            table = table_map.get(child)
            if table is None:
                continue

            table_md = _table_to_markdown(table)
            if table_md:
                # Tabloyu mevcut bloğa ekle (ardından flush edilecek)
                current_block_lines.append(table_md)

    # Son bloğu kapat
    _flush_block()

    if not chunks:
        return [{
            "id":   f"empty-{uuid.uuid4()}",
            "text": f"[{file_basename}] Dosya boş veya içerik bulunamadı.",
            "metadata": {"source": file_basename, "type": "text_empty"}
        }]

    # prev_id / next_id bağlantıları
    for i, chunk in enumerate(chunks):
        chunk["metadata"]["prev_id"] = chunks[i - 1]["id"] if i > 0 else ""
        chunk["metadata"]["next_id"] = chunks[i + 1]["id"] if i < len(chunks) - 1 else ""

    return chunks


def _table_to_markdown(table) -> str:
    """Word tablosunu Markdown formatına dönüştürür."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Header separator
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows) if rows else ""


# ── Düz metin okuyucu ─────────────────────────────────────────────────

def _read_plain(file_path: str) -> str:
    """TXT / MD → ham metin"""
    encodings = ["utf-8-sig", "utf-8", "latin-1", "cp1254"]
    for enc in encodings:
        try:
            with open(file_path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(file_path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _chunks_from_text(raw_text: str, file_basename: str, ext: str) -> list[dict]:
    """Düz metin için klasik parçalama + belge adı ön eki."""
    if not raw_text.strip():
        return [{
            "id":   f"empty-{uuid.uuid4()}",
            "text": f"[{file_basename}] Dosya boş veya içerik bulunamadı.",
            "metadata": {"source": file_basename, "type": "text_empty"}
        }]

    sub_chunks = _chunk_text(raw_text)
    total = len(sub_chunks)
    ids = [str(uuid.uuid4()) for _ in sub_chunks]
    chunks = []
    for idx, part in enumerate(sub_chunks):
        chunks.append({
            "id":   ids[idx],
            "text": f"[{file_basename} | Parça {idx+1}/{total}]\n\n{part}",
            "metadata": {
                "page":        idx + 1,
                "chunk_index": idx + 1,
                "source":      file_basename,
                "type":        f"text_{ext}",
                "total_pages": total,
                "prev_id":     ids[idx - 1] if idx > 0 else "",
                "next_id":     ids[idx + 1] if idx < total - 1 else "",
            }
        })
    return chunks


def _chunk_text(text: str) -> list[str]:
    """Paragraf sınırını koruyarak chunk'lara böl."""
    chunks = []
    start  = 0
    length = len(text)

    while start < length:
        end = start + _CHUNK_SIZE
        if end < length:
            # İleri doğru (son kısımdan) güvenli bir bitiş noktası ara
            for sep in ("\n\n", "\n", ". ", " "):
                pos = text.rfind(sep, start, end)
                if pos != -1 and pos > start:
                    end = pos + len(sep)
                    break
                    
        part = text[start:end].strip()
        if part:
            chunks.append(part)
            
        # Kesişim (Overlap) ayarı:
        raw_start = end - _CHUNK_OVERLAP
        if raw_start >= length or raw_start <= start:
            break
            
        # Geriye doğru düz saymak yerine, bulunduğumuz yere en yakın anlamlı başa git
        new_start = raw_start
        for sep in ("\n\n", "\n", ". ", " "):
            pos = text.find(sep, raw_start, end)
            if pos != -1:
                new_start = pos + len(sep)
                break
                
        # Eğer new_start end'e takılı kaldıysa (sonsuz döngüyü önle)
        if new_start >= end or new_start <= start:
            start = end
        else:
            start = new_start

    return chunks
